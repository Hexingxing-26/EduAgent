import os
from typing import List


class Document:
    def __init__(self, content: str, metadata: dict = None):
        self.content = content
        self.metadata = metadata or {}

    def to_dict(self):
        return {"content": self.content, "metadata": self.metadata}


class DocumentLoader:
    def __init__(self, data_dir: str = "data"):
        self.data_dir = data_dir
        self.knowledge_base_dir = os.path.join("数据", "数据", "知识库数据")

    def load_file(self, file_path: str) -> Document:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        try:
            if ext in [".txt", ".md", ".json"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            elif ext in [".pdf"]:
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(file_path)
                    content = "\n".join([page.extract_text() for page in reader.pages])
                except ImportError:
                    raise ImportError("请安装 PyPDF2: pip install PyPDF2")
            elif ext in [".docx"]:
                try:
                    from docx import Document as DocxDoc
                    docx_doc = DocxDoc(file_path)
                    content = "\n".join([paragraph.text for paragraph in docx_doc.paragraphs])
                except ImportError:
                    raise ImportError("请安装 python-docx: pip install python-docx")
            else:
                raise ValueError(f"不支持的文件格式: {ext}")

            return Document(
                content=content,
                metadata={
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_type": ext
                }
            )
        except Exception as e:
            print(f">>> 加载文件失败 {file_path}: {e}")
            return Document(content="", metadata={"file_path": file_path, "error": str(e)})

    def load_knowledge_base(self) -> List[Document]:
        documents = []
        if not os.path.exists(self.knowledge_base_dir):
            print(f">>> 知识库目录不存在: {self.knowledge_base_dir}")
            return documents

        supported_extensions = [".txt", ".md", ".pdf", ".docx"]
        total_files = 0
        loaded_files = 0

        for root, dirs, files in os.walk(self.knowledge_base_dir):
            for filename in files:
                file_path = os.path.join(root, filename)
                _, ext = os.path.splitext(filename)
                if ext.lower() in supported_extensions:
                    total_files += 1
                    doc = self.load_file(file_path)
                    if doc.content:
                        relative_path = os.path.relpath(file_path, self.knowledge_base_dir)
                        print(f">>> 加载知识库文件: {relative_path}, 内容长度: {len(doc.content)}")
                        doc.metadata["relative_path"] = relative_path
                        documents.append(doc)
                        loaded_files += 1

        print(f">>> 知识库加载完成: 共发现 {total_files} 个文件, 成功加载 {loaded_files} 个")
        return documents

    def load_directory(self) -> List[Document]:
        documents = []
        
        if os.path.exists(self.data_dir):
            supported_extensions = [".txt", ".md", ".pdf"]
            for filename in os.listdir(self.data_dir):
                file_path = os.path.join(self.data_dir, filename)
                _, ext = os.path.splitext(filename)
                if os.path.isfile(file_path) and ext.lower() in supported_extensions:
                    doc = self.load_file(file_path)
                    if doc.content:
                        documents.append(doc)

        kb_docs = self.load_knowledge_base()
        documents.extend(kb_docs)

        return documents


class TextSplitter:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        chunks = []
        text_length = len(text)

        if text_length <= self.chunk_size:
            return [text.strip()] if text.strip() else []

        start = 0
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            chunk = text[start:end]

            if end < text_length:
                last_period = chunk.rfind("。")
                last_newline = chunk.rfind("\n")
                split_pos = max(last_period, last_newline)
                if split_pos > self.chunk_size - self.chunk_overlap:
                    end = start + split_pos + 1
                    chunk = text[start:end]

            chunks.append(chunk.strip())
            
            if end >= text_length:
                break
                
            start = end - self.chunk_overlap
            
            if start <= 0:
                start = end

        return chunks

    def split_document(self, doc: Document) -> List[Document]:
        chunks = self.split_text(doc.content)
        return [
            Document(
                content=chunk,
                metadata={
                    **doc.metadata,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            )
            for i, chunk in enumerate(chunks)
        ]

    def split_documents(self, docs: List[Document]) -> List[Document]:
        all_chunks = []
        for doc in docs:
            chunks = self.split_document(doc)
            all_chunks.extend(chunks)
        return all_chunks